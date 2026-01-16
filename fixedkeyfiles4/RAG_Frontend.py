import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
import sys
import os
import time
import csv
import json
import pandas as pd

try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

# å¼•å…¥å¿…è¦çš„ PyQt5 ç»„ä»¶ (æ–°å¢ QDialog, QListWidget ç­‰)
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, 
                             QMessageBox, QSplitter, QComboBox, QCheckBox, QRadioButton, 
                             QButtonGroup, QFrame, QGroupBox, QInputDialog, QDialog, QListWidget)

from PyQt5.QtCore import Qt, QSettings
from PyQt5.QtGui import QTextCursor, QColor

# å¼•å…¥ Backend é€»è¾‘
from RAG_Backend import RecallWorker

# ================= é»˜è®¤èˆªå¸åˆ—è¡¨ (ä¸‰æ­¥èµ°ç­–ç•¥ - Step 1 æ•°æ®æº) =================
DEFAULT_AIRLINES = [
    "ä¸­å›½å›½é™…èˆªç©º", "å—æ–¹èˆªç©º", "ä¸œæ–¹èˆªç©º", "æµ·å—èˆªç©º",
    "å¦é—¨èˆªç©º", "å››å·èˆªç©º", "æ·±åœ³èˆªç©º", "æ˜¥ç§‹èˆªç©º",
    "å‰ç¥¥èˆªç©º", "é¦–éƒ½èˆªç©º", "å±±ä¸œèˆªç©º", "å¤©æ´¥èˆªç©º",
    "ä¸Šæµ·èˆªç©º", "ç¥¥é¹èˆªç©º", "è¥¿éƒ¨èˆªç©º", "é•¿é¾™èˆªç©º",
    "Air China", "China Southern", "China Eastern"
]
AIRLINE_DICT_FILE = "airline_dict.txt"

# ================= æ ·å¼è¡¨ (Dark Mode) =================
STYLESHEET = """
QMainWindow { background-color: #2b2b2b; color: #e0e0e0; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
QLabel { color: #aaaaaa; font-weight: bold; font-size: 13px; }
QLineEdit { background-color: #3c3c3c; color: #ffffff; border: 1px solid #555555; padding: 6px; border-radius: 4px; }
QTextEdit { background-color: #1e1e1e; color: #e0e0e0; border: 1px solid #444444; font-family: Consolas, monospace; font-size: 12px; }
QPushButton { background-color: #007acc; color: white; border: none; padding: 8px 16px; border-radius: 4px; font-weight: bold; font-size: 14px; }
QPushButton:hover { background-color: #005f9e; }
QPushButton:pressed { background-color: #004a80; }
QPushButton:disabled { background-color: #444444; color: #888888; }
/* Stop Button Style */
QPushButton#StopBtn { background-color: #d32f2f; }
QPushButton#StopBtn:hover { background-color: #b71c1c; }
QComboBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QComboBox::drop-down { border: 0px; }
QRadioButton { color: #e0e0e0; font-weight: bold; spacing: 5px; }
QRadioButton::indicator { width: 16px; height: 16px; }
QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; font-weight: bold; color: #aaa; }
QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; padding: 0 5px; }
QFrame#Divider { border: 1px solid #444444; }
/* æ¨¡å¼é€‰æ‹©ç‰¹å®šæ ·å¼ */
QRadioButton#ModeSmart { color: #4facfe; }
QRadioButton#ModePrecise { color: #00f260; }
QRadioButton#ModeFuzzy { color: #ff9a9e; }
/* ListWidget for Dict Editor */
QListWidget { background-color: #333; color: white; border: 1px solid #555; }
"""

# ================= èˆªå¸å­—å…¸ç¼–è¾‘å™¨çª—å£ =================
class AirlineDictEditor(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("ç®¡ç†èˆªå¸å­—å…¸ (airline_dict.txt)")
        self.resize(400, 500)
        self.setStyleSheet(STYLESHEET)
        self.init_ui()
        self.load_data()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # è¾“å…¥åŒº
        input_layout = QHBoxLayout()
        self.entry_new = QLineEdit()
        self.entry_new.setPlaceholderText("è¾“å…¥æ–°èˆªå¸åç§°...")
        btn_add = QPushButton("â• æ·»åŠ ")
        btn_add.clicked.connect(self.add_item)
        btn_add.setFixedWidth(80)
        input_layout.addWidget(self.entry_new)
        input_layout.addWidget(btn_add)
        layout.addLayout(input_layout)

        # åˆ—è¡¨åŒº
        self.list_widget = QListWidget()
        layout.addWidget(self.list_widget)

        # åº•éƒ¨æŒ‰é’®
        btn_layout = QHBoxLayout()
        btn_del = QPushButton("âŒ åˆ é™¤é€‰ä¸­")
        btn_del.setStyleSheet("background-color: #c0392b;")
        btn_del.clicked.connect(self.delete_item)
        
        btn_save = QPushButton("ğŸ’¾ ä¿å­˜å¹¶å…³é—­")
        btn_save.setStyleSheet("background-color: #27ae60;")
        btn_save.clicked.connect(self.save_and_close)
        
        btn_layout.addWidget(btn_del)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_save)
        layout.addLayout(btn_layout)

    def load_data(self):
        # ç¡®ä¿æ–‡ä»¶å­˜åœ¨
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except Exception:
                pass
        
        # è¯»å–
        try:
            with open(AIRLINE_DICT_FILE, "r", encoding="utf-8") as f:
                lines = f.readlines()
                for line in lines:
                    text = line.strip()
                    if text:
                        self.list_widget.addItem(text)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"è¯»å–å­—å…¸å¤±è´¥: {str(e)}")

    def add_item(self):
        text = self.entry_new.text().strip()
        if not text:
            return
        # æŸ¥é‡
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        if text in items:
            QMessageBox.warning(self, "æç¤º", "è¯¥å…³é”®è¯å·²å­˜åœ¨")
            return
        
        self.list_widget.addItem(text)
        self.entry_new.clear()

    def delete_item(self):
        row = self.list_widget.currentRow()
        if row >= 0:
            self.list_widget.takeItem(row)

    def save_and_close(self):
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        try:
            with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                for item in items:
                    f.write(item + "\n")
            QMessageBox.information(self, "æˆåŠŸ", "èˆªå¸å­—å…¸å·²æ›´æ–°ï¼")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "é”™è¯¯", f"ä¿å­˜å¤±è´¥: {str(e)}")


# ================= ä¸»ç•Œé¢ =================
class RAGRecallApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("RAG å·¥ä¸šçº§å…¨æµç¨‹ (Recall + RRF Fusion + Multi-Model Support)")
        self.resize(1400, 950) 
        self.setStyleSheet(STYLESHEET)
        
        self.settings = QSettings("MyCorp", "RAGRecall_Final_v9_MultiModel")
        self.cached_results = []
        self.cached_summary = ""
        self.cached_query = ""
        self.worker = None 
        
        # åˆå§‹åŒ–å­—å…¸æ–‡ä»¶
        self.ensure_airline_dict()
        
        self.init_ui()
        
    def ensure_airline_dict(self):
        """ç¡®ä¿èˆªå¸å­—å…¸æ–‡ä»¶å­˜åœ¨"""
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except:
                pass

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(10)
        
        # === Left Widget ===
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # 1. DB & JSON Inputs
        db_layout = QHBoxLayout()
        self.db_path_edit = QLineEdit()
        self.db_path_edit.setText(self.settings.value("last_db_path", ""))
        btn_db = QPushButton("ğŸ“‚ å‘é‡åº“")
        btn_db.clicked.connect(self.browse_db)
        db_layout.addWidget(QLabel("Vector DB:"))
        db_layout.addWidget(self.db_path_edit)
        db_layout.addWidget(btn_db)
        left_layout.addLayout(db_layout)
        
        json_layout = QHBoxLayout()
        self.json_path_edit = QLineEdit()
        self.json_path_edit.setText(self.settings.value("last_json_path", ""))
        btn_json = QPushButton("ğŸ“„ PageIndex")
        btn_json.clicked.connect(self.browse_json)
        json_layout.addWidget(QLabel("Structure JSON:"))
        json_layout.addWidget(self.json_path_edit)
        json_layout.addWidget(btn_json)
        left_layout.addLayout(json_layout)
        
        # 2. Search Mode & Model Selection
        mode_layout = QHBoxLayout()
        
        # 2.1 ç­–ç•¥é€‰æ‹©
        mode_layout.addWidget(QLabel("ğŸ” ç­–ç•¥:"))
        self.mode_group = QButtonGroup(self)
        
        self.radio_smart = QRadioButton("ğŸ”µ æ™ºèƒ½")
        self.radio_smart.setObjectName("ModeSmart")
        self.radio_smart.setChecked(True)
        self.mode_group.addButton(self.radio_smart, 1)
        mode_layout.addWidget(self.radio_smart)
        
        self.radio_precise = QRadioButton("ğŸŸ¢ ç²¾å‡†")
        self.radio_precise.setObjectName("ModePrecise")
        self.mode_group.addButton(self.radio_precise, 2)
        mode_layout.addWidget(self.radio_precise)
        
        self.radio_fuzzy = QRadioButton("ğŸŸ¡ æ¨¡ç³Š")
        self.radio_fuzzy.setObjectName("ModeFuzzy")
        self.mode_group.addButton(self.radio_fuzzy, 3)
        mode_layout.addWidget(self.radio_fuzzy)
        
        mode_layout.addSpacing(15)

        # 2.2 Doc Type Selection & Dynamic Throttling
        mode_layout.addWidget(QLabel("ğŸ“‚ æ–‡æ¡£åå¥½:"))
        self.combo_doc_type = QComboBox()
        self.combo_doc_type.addItems([
            "ä¸æŒ‡å®šç±»å‹",
            "æ•°æ®è¡¨",
            "å…¬å¸å…¬æ–‡",
            "ä¹¦ç±/æ•™æ",   # æ–°å¢
            "é•¿ç¯‡è®ºæ–‡",     # æ–°å¢
            "æŠ€æœ¯æ–‡æ¡£",
            "æ³•å¾‹æ¡æ–‡",
            "LLM OCRæ–‡æ¡£",
            "LLMç”Ÿæˆæ€»ç»“"
        ])
        self.combo_doc_type.setFixedWidth(120)
        self.combo_doc_type.currentTextChanged.connect(self.on_doc_type_changed)
        mode_layout.addWidget(self.combo_doc_type)

        # === æ–°å¢: ä¹¦ç±æ¨¡å¼ä¸‹çš„æµæ§ UI (é»˜è®¤éšè—) ===
        self.chunk_limit_widget = QWidget()
        self.chunk_limit_layout = QHBoxLayout(self.chunk_limit_widget)
        self.chunk_limit_layout.setContentsMargins(0, 0, 0, 0)
        
        self.lbl_chunk = QLabel("åˆ‡å—é™åˆ¶:")
        self.lbl_chunk.setStyleSheet("color: #ff9f43;") # æ©™è‰²è­¦ç¤º
        self.combo_chunk_limit = QComboBox()
        self.combo_chunk_limit.addItems([
            "15 (æé€Ÿ)", 
            "25 (å¹³è¡¡)", 
            "40 (æ·±åº¦/é£é™©)"
        ])
        self.combo_chunk_limit.setFixedWidth(100)
        self.combo_chunk_limit.setCurrentIndex(1) # é»˜è®¤ 25
        
        self.chunk_limit_layout.addWidget(self.lbl_chunk)
        self.chunk_limit_layout.addWidget(self.combo_chunk_limit)
        
        mode_layout.addWidget(self.chunk_limit_widget)
        self.chunk_limit_widget.hide() # åˆå§‹éšè—

        mode_layout.addSpacing(15)

        # 2.3 Model Selection
        mode_layout.addWidget(QLabel("ğŸ§  æ¨¡å‹:"))
        self.combo_model = QComboBox()
        self.combo_model.addItems([
            "DeepSeek-R1", 
            "DeepSeek-V3", 
            "X1-70B-thinking", 
            "X1-70B-fast"
        ])
        self.combo_model.setFixedWidth(140)
        mode_layout.addWidget(self.combo_model)
        
        mode_layout.addStretch()
        left_layout.addLayout(mode_layout)

        # 3. Knowledge Config (Stopwords + Airline Dict)
        config_group = QGroupBox("ğŸ”§ çŸ¥è¯†åº“é…ç½® (å­—å…¸/åœç”¨è¯)")
        config_layout = QVBoxLayout(config_group)
        
        # 3.1 Input Area
        self.stopwords_edit = QTextEdit()
        self.stopwords_edit.setPlaceholderText("åœ¨æ­¤è¾“å…¥åœç”¨è¯ï¼Œä»¥é€—å·åˆ†éš”...")
        self.stopwords_edit.setMaximumHeight(50)
        config_layout.addWidget(self.stopwords_edit)
        
        # 3.2 Control Buttons
        sw_btn_layout = QHBoxLayout()
        
        self.sw_name_edit = QLineEdit()
        self.sw_name_edit.setPlaceholderText("Config Name")
        self.sw_name_edit.setFixedWidth(100)
        sw_btn_layout.addWidget(self.sw_name_edit)
        
        btn_sw_save = QPushButton("ğŸ’¾ Save Stopwords")
        btn_sw_save.clicked.connect(self.save_stopwords)
        btn_sw_save.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_save)
        
        btn_sw_load = QPushButton("ğŸ“‚ Import")
        btn_sw_load.clicked.connect(self.import_stopwords)
        btn_sw_load.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_load)
        
        # æ–°å¢ï¼šèˆªå¸å­—å…¸æŒ‰é’®
        btn_airline = QPushButton("âœˆï¸ ç®¡ç†èˆªå¸å­—å…¸")
        btn_airline.clicked.connect(self.open_airline_editor)
        btn_airline.setFixedHeight(28)
        btn_airline.setStyleSheet("background-color: #8e44ad; color: white;")
        sw_btn_layout.addWidget(btn_airline)
        
        sw_btn_layout.addStretch()
        config_layout.addLayout(sw_btn_layout)
        
        left_layout.addWidget(config_group)

        # 4. Query Input
        left_layout.addWidget(QLabel("ç”¨æˆ·æŸ¥è¯¢ (Query):"))
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("è¯·è¾“å…¥é—®é¢˜ (å¦‚ï¼šæŸ¥æ‰¾èˆªç­ JMUï¼Œæˆ–è¯¢é—®æŸæ“ä½œæµç¨‹)...")
        self.query_input.setMaximumHeight(60)
        left_layout.addWidget(self.query_input)
        
        # 5. Search & Stop Buttons
        btn_layout = QHBoxLayout()
        
        self.btn_search = QPushButton("ğŸš€ æ‰§è¡Œå…¨æµç¨‹ (Recall -> RRF -> Summary)")
        self.btn_search.setFixedHeight(45)
        self.btn_search.setStyleSheet("background-color: #2da44e; font-size: 15px;")
        self.btn_search.clicked.connect(self.start_recall)
        btn_layout.addWidget(self.btn_search)
        
        self.btn_stop = QPushButton("ğŸ›‘ åœæ­¢è¿è¡Œ")
        self.btn_stop.setObjectName("StopBtn")
        self.btn_stop.setFixedHeight(45)
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.stop_recall)
        btn_layout.addWidget(self.btn_stop)
        
        left_layout.addLayout(btn_layout)
        
        # 6. Result Displays
        left_layout.addWidget(QLabel("ğŸ¤– æ™ºèƒ½æ€»ç»“ (Thinking + Answer):"))
        self.summary_display = QTextEdit()
        self.summary_display.setReadOnly(True)
        self.summary_display.setStyleSheet("""
            QTextEdit {
                background-color: #252526; 
                color: #dcdcaa; 
                font-family: 'Segoe UI', sans-serif; 
                font-size: 14px; 
                border: 1px solid #007acc;
                line-height: 1.6;
            }
        """)
        self.summary_display.setMinimumHeight(250)
        left_layout.addWidget(self.summary_display)

        left_layout.addWidget(QLabel("ğŸ“š RRF Fused Context (Top-12 Unique):"))
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setStyleSheet("font-family: Consolas; font-size: 12px; color: #aaddff;")
        left_layout.addWidget(self.result_display)
        
        # 7. Export Area
        export_layout = QHBoxLayout()
        export_layout.addWidget(QLabel("å¯¼å‡ºæ ¼å¼:"))
        
        self.combo_format = QComboBox()
        self.combo_format.addItems(["xlsx", "csv", "txt", "docx", "md"])
        self.combo_format.setFixedWidth(100)
        export_layout.addWidget(self.combo_format)
        
        self.btn_export = QPushButton("ğŸ’¾ å¯¼å‡ºç»“æœ")
        self.btn_export.setStyleSheet("background-color: #d2691e;")
        self.btn_export.clicked.connect(self.export_data)
        export_layout.addWidget(self.btn_export)
        
        export_layout.addStretch() 
        left_layout.addLayout(export_layout)

        # === Right Console ===
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("ğŸ“Ÿ System Console"))
        self.console_output = QTextEdit()
        self.console_output.setReadOnly(True)
        self.console_output.setStyleSheet("background-color: #111; color: #0f0; font-family: Consolas;")
        right_layout.addWidget(self.console_output)
        
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([900, 400])
        main_layout.addWidget(splitter)
        
        # åŠ è½½ä¸Šæ¬¡çš„ Stopwords
        last_sw = self.settings.value("last_stopwords", "")
        if last_sw:
            self.stopwords_edit.setText(last_sw)

    # ================= UI äº¤äº’é€»è¾‘ =================
    
    def on_doc_type_changed(self, text):
        """åŠ¨æ€æ˜¾ç¤ºæµæ§ç»„ä»¶"""
        if text in ["ä¹¦ç±/æ•™æ", "é•¿ç¯‡è®ºæ–‡"]:
            self.chunk_limit_widget.show()
            self.log(f"ğŸ“š æ£€æµ‹åˆ°ä¹¦ç±æ¨¡å¼ï¼Œå·²å¯ç”¨åŠ¨æ€æµæ§ (é»˜è®¤ 25 chunks)")
        else:
            self.chunk_limit_widget.hide()

    def browse_db(self):
        path, _ = QFileDialog.getOpenFileName(self, "é€‰æ‹©æ•°æ®åº“", "", "SQLite DB (*.db);;All Files (*.*)")
        if path:
            self.db_path_edit.setText(path)
            self.settings.setValue("last_db_path", path)

    def browse_json(self):
        path, _ = QFileDialog.getOpenFileName(self, "é€‰æ‹© JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            self.json_path_edit.setText(path)
            self.settings.setValue("last_json_path", path)

    def log(self, msg):
        self.console_output.append(msg)
        cursor = self.console_output.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.console_output.setTextCursor(cursor)

    def update_summary(self, text):
        """æµå¼æ›´æ–°æ€»ç»“æ–‡æœ¬"""
        self.cached_summary = text 
        self.summary_display.setMarkdown(text) 
        cursor = self.summary_display.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.summary_display.setTextCursor(cursor)

    def display_results(self, results):
        self.cached_results = results 
        html = ""
        for item in results:
            score_text = f"{item['final_score']:.4f}"
            debug_info = item.get('debug_score', '')
            source = item.get('source', 'VECTOR')
            
            if "JSON" in source:
                source_style = "background-color: #00f260; color: #111; padding: 2px 5px; border-radius: 3px; font-weight: bold;"
            else:
                source_style = "background-color: #007acc; color: white; padding: 2px 5px; border-radius: 3px;"
            
            score_color = "#00ff00" 
            
            html += f"""
            <div style='border-bottom: 1px solid #555; padding: 12px; margin-bottom: 8px;'>
                <span style='color: #888; font-weight:bold;'>Rank #{item['rank']}</span> | 
                <span style='{source_style}'>{source}</span> | 
                <span style='color: {score_color}; font-weight: bold;'>RRF Score: {score_text}</span> 
                <span style='color: #aaa; font-size:11px;'>[{debug_info}]</span><br>
                <div style='margin-top:5px; color: #ffcc00;'><b>[Section Path]</b> {item['path']}</div>
                <div style='margin-top:5px; background-color: #222; padding: 8px; border-left: 3px solid #2da44e; white-space: pre-wrap;'>
{item['content'][:200]}...
                </div>
            </div>
            """
        self.result_display.setHtml(html)

    # ================= çŸ¥è¯†åº“é…ç½®åŠŸèƒ½ (Stopwords + Airline) =================
    
    def open_airline_editor(self):
        """æ‰“å¼€èˆªå¸å­—å…¸ç¼–è¾‘å™¨"""
        editor = AirlineDictEditor(self)
        editor.exec_() # æ¨¡æ€è¿è¡Œ

    def get_current_stopwords(self):
        text = self.stopwords_edit.toPlainText().strip()
        if not text:
            return []
        return [w.strip() for w in text.replace('ï¼Œ', ',').split(',') if w.strip()]
    
    def get_airline_list(self):
        """è¯»å–èˆªå¸å­—å…¸"""
        airlines = []
        if os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, 'r', encoding='utf-8') as f:
                    airlines = [line.strip() for line in f if line.strip()]
            except:
                pass
        return airlines

    def save_stopwords(self):
        name = self.sw_name_edit.text().strip()
        if not name:
            QMessageBox.warning(self, "Warning", "Please enter a profile name.")
            return
        
        sw_list = self.get_current_stopwords()
        filename = f"stopwords_{name}.json"
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump({"name": name, "stopwords": sw_list}, f, ensure_ascii=False, indent=2)
            self.log(f"ğŸ’¾ Stopwords profile '{name}' saved to {filename}")
            QMessageBox.information(self, "Success", f"Saved to {filename}")
        except Exception as e:
            self.log(f"âŒ Save failed: {str(e)}")

    def import_stopwords(self):
        path, _ = QFileDialog.getOpenFileName(self, "Import Stopwords JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                if isinstance(data, list):
                    sw_list = data
                elif isinstance(data, dict) and "stopwords" in data:
                    sw_list = data["stopwords"]
                else:
                    raise ValueError("Invalid JSON format")
                
                self.stopwords_edit.setText(", ".join(sw_list))
                self.log(f"ğŸ“‚ Imported {len(sw_list)} words from {os.path.basename(path)}")
            except Exception as e:
                self.log(f"âŒ Import failed: {str(e)}")

    # ================= æ ¸å¿ƒæµç¨‹æ§åˆ¶ =================

    def start_recall(self):
        # 1. åŸºç¡€æ ¡éªŒ
        db_path = self.db_path_edit.text().strip()
        json_path = self.json_path_edit.text().strip()
        query = self.query_input.toPlainText().strip()
        
        if not db_path or not os.path.exists(db_path):
            QMessageBox.warning(self, "Error", "æ— æ•ˆçš„æ•°æ®åº“è·¯å¾„")
            return
        if not query:
            return

        # 2. è·å–å‚æ•°
        try:
            mode_id = self.mode_group.checkedId()
            search_mode = "smart" 
            if mode_id == 2: search_mode = "precise"
            elif mode_id == 3: search_mode = "fuzzy"

            summary_model = self.combo_model.currentText()
            doc_type = self.combo_doc_type.currentText()
            stopwords = self.get_current_stopwords()
            airline_names = self.get_airline_list()

            # 3. å¤„ç†åˆ‡å—é™åˆ¶ (æµæ§é€»è¾‘)
            chunk_limit = 40 # é»˜è®¤å€¼
            if self.chunk_limit_widget.isVisible():
                # ä» "25 (å¹³è¡¡)" è¿™æ ·çš„å­—ç¬¦ä¸²ä¸­æå–æ•°å­—
                raw_text = self.combo_chunk_limit.currentText()
                try:
                    chunk_limit = int(raw_text.split()[0])
                except:
                    chunk_limit = 25
            
            self.cached_query = query
            self.cached_results = []
            self.cached_summary = ""
            
            # UI çŠ¶æ€æ›´æ–°
            self.btn_search.setEnabled(False)
            self.btn_stop.setEnabled(True) 
            self.result_display.clear()
            self.summary_display.clear() 
            self.console_output.clear()
            
            mode_text = {"smart": "ğŸ”µ æ™ºèƒ½èåˆ", "precise": "ğŸŸ¢ ç²¾å‡†æŸ¥è¡¨", "fuzzy": "ğŸŸ¡ æ¨¡ç³Šå’¨è¯¢"}[search_mode]
            self.log(f"ğŸš€ åˆå§‹åŒ–ä»»åŠ¡... | ç­–ç•¥: {mode_text} | æ¨¡å‹: {summary_model}")
            self.log(f"â„¹ï¸ æ–‡æ¡£ç±»å‹: {doc_type} | Limit: {chunk_limit}")
            self.log(f"â„¹ï¸ çŸ¥è¯†åº“: èˆªå¸è¯è¡¨ {len(airline_names)} | Stopwords {len(stopwords)}")
            
            # å®ä¾‹åŒ–åç«¯ Worker
            self.worker = RecallWorker(
                query_text=query, 
                db_path=db_path, 
                json_path=json_path, 
                search_mode=search_mode, 
                summary_model=summary_model, 
                doc_type=doc_type, 
                stopwords=stopwords, 
                airline_names=airline_names,
                chunk_limit=chunk_limit # ä¼ å…¥æµæ§å‚æ•°
            )
            
            self.worker.log_signal.connect(self.log)
            self.worker.result_signal.connect(self.display_results)
            self.worker.summary_signal.connect(self.update_summary) 
            self.worker.finish_signal.connect(self.on_finished)
            self.worker.start()

        except Exception as e:
            # å…œåº•ï¼šé˜²æ­¢å‰ç«¯é€»è¾‘å´©æºƒ
            QMessageBox.critical(self, "System Error", f"å¯åŠ¨å¤±è´¥: {str(e)}")
            self.btn_search.setEnabled(True)
            self.btn_stop.setEnabled(False)

    def stop_recall(self):
        if self.worker and self.worker.isRunning():
            self.log("ğŸ›‘ ç”¨æˆ·ç‚¹å‡»åœæ­¢ï¼Œæ­£åœ¨è¯·æ±‚ Backend ä¸­æ–­...")
            self.worker.stop()
            self.btn_stop.setEnabled(False) 
            self.btn_stop.setText("Stopping...")

    def on_finished(self, success):
        self.btn_search.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.btn_stop.setText("ğŸ›‘ åœæ­¢è¿è¡Œ")
        
        if success:
            self.log("âœ… å…¨æµç¨‹ç»“æŸ")
        else:
            self.log("âŒ æµç¨‹è¢«ä¸­æ–­æˆ–å‘ç”Ÿé”™è¯¯")
            if not self.cached_summary:
                self.summary_display.setText("[ Process Stopped / Failed ]")

    # ================= å¯¼å‡ºåŠŸèƒ½ =================
    def export_data(self):
        if not self.cached_summary and not self.cached_results:
            QMessageBox.warning(self, "æç¤º", "å½“å‰æ²¡æœ‰å¯å¯¼å‡ºçš„ç»“æœï¼Œè¯·å…ˆæ‰§è¡ŒæŸ¥è¯¢ã€‚")
            return

        fmt = self.combo_format.currentText()
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"export_{timestamp}.{fmt}"
        save_path = os.path.join(os.getcwd(), filename)

        try:
            self.log(f"ğŸ’¾ æ­£åœ¨å¯¼å‡ºä¸º {fmt} ...")
            
            if fmt == "xlsx":
                if pd is None:
                    raise ImportError("ç¼ºå°‘ pandas æˆ– openpyxl åº“")
                
                data_rows = []
                for item in self.cached_results:
                    data_rows.append({
                        "Rank": item['rank'],
                        "Source": item.get('source', 'VECTOR'),
                        "RRF Score": item['final_score'],
                        "Debug Score": item.get('debug_score', ''),
                        "Section Path": item['path'],
                        "Content": item['content']
                    })
                
                df = pd.DataFrame(data_rows)
                
                with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                    summary_df = pd.DataFrame([["DeepSeek R1 Summary"], [self.cached_summary], [""]])
                    summary_df.to_excel(writer, sheet_name='Report', index=False, header=False, startrow=0)
                    
                    pd.DataFrame([["Top Results"]]).to_excel(writer, sheet_name='Report', index=False, header=False, startrow=4)
                    df.to_excel(writer, sheet_name='Report', index=False, startrow=6)
            
            elif fmt == "csv":
                with open(save_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    writer.writerow(["=== DeepSeek R1 Summary ==="])
                    writer.writerow([self.cached_summary])
                    writer.writerow([])
                    writer.writerow(["=== Top Results ==="])
                    writer.writerow(["Rank", "Source", "RRF Score", "Debug Score", "Section Path", "Content"])
                    for item in self.cached_results:
                        writer.writerow([
                            item['rank'],
                            item.get('source', 'VECTOR'),
                            f"{item['final_score']:.4f}",
                            item.get('debug_score', ''),
                            item['path'],
                            item['content']
                        ])

            elif fmt == "txt":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"Query: {self.cached_query}\n")
                    f.write("="*50 + "\n")
                    f.write("DeepSeek R1 Summary:\n")
                    f.write("="*50 + "\n")
                    clean_summary = self.cached_summary.replace("**", "").replace(">", "")
                    f.write(clean_summary + "\n\n")
                    f.write("="*50 + "\n")
                    f.write("Top Results:\n")
                    f.write("="*50 + "\n")
                    for item in self.cached_results:
                        f.write(f"[Rank #{item['rank']}] [{item.get('source','VECTOR')}] RRF: {item['final_score']:.4f}\n")
                        f.write(f"Path: {item['path']}\n")
                        f.write(f"Content:\n{item['content']}\n")
                        f.write("-" * 30 + "\n")

            elif fmt == "md":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"# RAG Query Report\n\n")
                    f.write(f"**Query:** {self.cached_query}\n\n")
                    f.write(f"## ğŸ¤– DeepSeek R1 Summary\n\n")
                    f.write(self.cached_summary + "\n\n")
                    f.write(f"## ğŸ“š Top Results\n\n")
                    for item in self.cached_results:
                        f.write(f"### Rank #{item['rank']} [{item.get('source','VECTOR')}] (RRF: {item['final_score']:.4f})\n")
                        f.write(f"**Path:** `{item['path']}`\n\n")
                        f.write(f"**Content:**\n\n")
                        content_block = item['content'].replace('\n', '\n> ')
                        f.write(f"> {content_block}\n\n")
                        f.write("---\n")

            elif fmt == "docx":
                if docx is None:
                    raise ImportError("ç¼ºå°‘ python-docx åº“")
                
                doc = docx.Document()
                doc.add_heading('RAG Analysis Report', 0)
                
                p = doc.add_paragraph()
                p.add_run('Query: ').bold = True
                p.add_run(self.cached_query)
                
                doc.add_heading('DeepSeek R1 Summary', level=1)
                doc.add_paragraph(self.cached_summary)
                
                doc.add_heading('Top Results', level=1)
                
                for item in self.cached_results:
                    p_header = doc.add_paragraph()
                    run = p_header.add_run(f"Rank #{item['rank']} | [{item.get('source','VECTOR')}] | RRF: {item['final_score']:.4f}")
                    run.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(0, 100, 0)
                    
                    p_path = doc.add_paragraph()
                    p_path.add_run("Path: ").bold = True
                    p_path.add_run(item['path']).italic = True
                    
                    p_content = doc.add_paragraph(item['content'])
                    p_content.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("-" * 40)

            self.log(f"âœ… å¯¼å‡ºæˆåŠŸ: {save_path}")
            QMessageBox.information(self, "æˆåŠŸ", f"æ–‡ä»¶å·²å¯¼å‡ºè‡³:\n{save_path}")

        except ImportError as ie:
            self.log(f"âŒ å¯¼å‡ºå¤±è´¥ (åº“ç¼ºå¤±): {str(ie)}")
            QMessageBox.critical(self, "é”™è¯¯", f"å¯¼å‡ºå¤±è´¥ï¼Œç¼ºå°‘å¿…è¦åº“:\n{str(ie)}")
        except Exception as e:
            self.log(f"âŒ å¯¼å‡ºå¼‚å¸¸: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "é”™è¯¯", f"å¯¼å‡ºè¿‡ç¨‹ä¸­å‘ç”Ÿé”™è¯¯:\n{str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RAGRecallApp()
    window.show()
    sys.exit(app.exec_())