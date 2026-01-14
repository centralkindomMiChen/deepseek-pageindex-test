import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

import sys
import json
import os
import html
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTextEdit, QListWidget,
    QListWidgetItem, QFileDialog, QSplitter, QMessageBox,
    QComboBox, QShortcut, QSlider
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QTextCursor, QTextCharFormat, QColor

# 兼容不同 PyQt5 版本的 QKeySequence 位置
try:
    from PyQt5.QtGui import QKeySequence
except ImportError:
    from PyQt5.QtWidgets import QKeySequence

# --- 可选依赖导入 ---
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


class PGIRecallWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PageIndex - 知识召回查询中心 (DeepSeek适配版)")
        self.resize(1400, 900)

        self.data = None
        self.all_nodes = []          # 扁平化存储所有节点
        self.last_loaded_path = None # 记录最后加载的文件路径，用于刷新

        self.init_ui()
        self.apply_styles()
        self.setup_shortcuts()
        
        # 初始化字体大小 (触发滑块默认值)
        self.change_font_size(self.slider_font.value())

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)

        # --- 顶部工具栏 ---
        top_bar = QHBoxLayout()

        self.btn_load = QPushButton("📂 加载索引JSON")
        self.btn_load.clicked.connect(self.load_json)

        self.btn_refresh = QPushButton("🔄 刷新")
        self.btn_refresh.setToolTip("重新加载当前文件并显示全部节点")
        self.btn_refresh.clicked.connect(self.refresh_current_file)

        self.edit_search = QLineEdit()
        self.edit_search.setPlaceholderText("🔍 输入关键词进行全局内容召回（标题/正文/摘要）...")
        self.edit_search.returnPressed.connect(self.search_content)

        self.btn_search = QPushButton("执行召回")
        self.btn_search.clicked.connect(self.search_content)

        # 导出功能
        self.combo_export = QComboBox()
        self.combo_export.addItems(["DOCX (Word)", "TXT (纯文本)", "CSV (表格)", "XLSX (Excel)"])
        self.combo_export.setFixedWidth(150)

        self.btn_export = QPushButton("💾 导出全部节点")
        self.btn_export.clicked.connect(self.export_all_nodes)

        top_bar.addWidget(self.btn_load)
        top_bar.addWidget(self.btn_refresh)
        top_bar.addWidget(self.edit_search, 4)
        top_bar.addWidget(self.btn_search)
        top_bar.addSpacing(30)
        top_bar.addWidget(QLabel("导出格式:"))
        top_bar.addWidget(self.combo_export)
        top_bar.addWidget(self.btn_export)

        layout.addLayout(top_bar)

        # --- 主内容区：Splitter 分割 ---
        splitter = QSplitter(Qt.Horizontal)

        # 左侧：结果列表
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.addWidget(QLabel("召回结果列表:"))
        self.list_results = QListWidget()
        self.list_results.itemClicked.connect(self.display_node_detail)
        left_layout.addWidget(self.list_results)

        # 右侧：详情预览 + 正文检索
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(8)

        right_layout.addWidget(QLabel("节点详情预览:"))

        # 标题与元信息区
        self.txt_header = QTextEdit()
        self.txt_header.setReadOnly(True)
        self.txt_header.setMaximumHeight(150)
        self.txt_header.setStyleSheet("border: none; background-color: #0d1117;") 
        right_layout.addWidget(self.txt_header)

        # 正文内检索栏
        search_bar = QHBoxLayout()
        search_bar.addWidget(QLabel("🔎 正文检索:"))
        self.edit_inner_search = QLineEdit()
        self.edit_inner_search.setPlaceholderText("在此输入关键词高亮正文内容 (支持 Ctrl+F)")
        self.edit_inner_search.textChanged.connect(self.highlight_text_in_detail)
        self.edit_inner_search.setStyleSheet("""
            background-color: #21262d; 
            border: 1px solid #30363d; 
            color: #ffd700; 
            font-weight: bold;
            padding: 6px;
        """)
        search_bar.addWidget(self.edit_inner_search)
        right_layout.addLayout(search_bar)

        # 正文内容区
        self.txt_detail = QTextEdit()
        self.txt_detail.setReadOnly(True)
        right_layout.addWidget(self.txt_detail)

        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        layout.addWidget(splitter, 1)

        # --- 底部：字体调节栏 ---
        font_bar = QHBoxLayout()
        font_bar.setContentsMargins(0, 5, 0, 0)
        
        lbl_font_icon = QLabel("🔠 字号调节:")
        lbl_font_icon.setStyleSheet("color: #c9d1d9; font-weight: normal;")
        
        self.slider_font = QSlider(Qt.Horizontal)
        self.slider_font.setRange(12, 40)
        self.slider_font.setValue(30)
        self.slider_font.setFixedWidth(200)
        self.slider_font.valueChanged.connect(self.change_font_size)
        
        self.lbl_font_val = QLabel("30px")
        self.lbl_font_val.setStyleSheet("color: #58a6ff; font-weight: bold; min-width: 40px;")

        font_bar.addStretch()
        font_bar.addWidget(lbl_font_icon)
        font_bar.addWidget(self.slider_font)
        font_bar.addWidget(self.lbl_font_val)
        
        layout.addLayout(font_bar)

    def apply_styles(self):
        self.setStyleSheet("""
            QMainWindow { background-color: #0d1117; }
            QLabel { 
                color: #58a6ff; 
                font-family: 'Segoe UI', 'Microsoft YaHei'; 
                font-weight: bold; 
                font-size: 14px; 
            }
            QLineEdit { 
                background-color: #161b22; 
                border: 1px solid #30363d; 
                border-radius: 6px; 
                color: #c9d1d9; 
                padding: 8px; 
                font-family: 'Microsoft YaHei'; 
            }
            QLineEdit:focus { border: 1px solid #58a6ff; }
            QPushButton { 
                background-color: #238636; 
                color: white; 
                border: none; 
                padding: 8px 15px; 
                border-radius: 6px; 
                font-weight: bold; 
            }
            QPushButton:hover { background-color: #2ea043; }
            QPushButton:pressed { background-color: #1a6329; }
            QComboBox {
                background-color: #161b22;
                color: #c9d1d9;
                border: 1px solid #30363d;
                padding: 6px;
                border-radius: 6px;
            }
            QListWidget { 
                background-color: #0d1117; 
                border: 1px solid #30363d; 
                border-radius: 6px;
                color: #c9d1d9; 
                padding: 5px;
            }
            QListWidget::item { padding: 8px; }
            QListWidget::item:selected { 
                background-color: #1f6feb; 
                border-radius: 6px; 
                color: white; 
            }
            QTextEdit { 
                background-color: #0d1117; 
                border: 1px solid #30363d; 
                border-radius: 6px;
                color: #c9d1d9; 
                line-height: 1.6; 
                padding: 12px;
                font-family: Consolas, 'Microsoft YaHei';
            }
            QSplitter::handle { background-color: #30363d; width: 6px; }
            QSlider::groove:horizontal {
                border: 1px solid #30363d;
                height: 6px;
                background: #161b22;
                margin: 2px 0;
                border-radius: 3px;
            }
            QSlider::handle:horizontal {
                background: #58a6ff;
                border: 1px solid #58a6ff;
                width: 14px;
                height: 14px;
                margin: -5px 0;
                border-radius: 7px;
            }
        """)

    def change_font_size(self, size):
        """动态调整主要内容区域的字体大小"""
        self.lbl_font_val.setText(f"{size}px")
        
        base_style_list = f"""
            QListWidget {{
                background-color: #0d1117; 
                border: 1px solid #30363d; 
                border-radius: 6px;
                color: #c9d1d9; 
                padding: 5px;
                font-size: {size}px;
            }}
        """
        
        base_style_text = f"""
            QTextEdit {{
                background-color: #0d1117; 
                border: 1px solid #30363d; 
                border-radius: 6px;
                color: #c9d1d9; 
                line-height: 1.6; 
                padding: 12px;
                font-family: Consolas, 'Microsoft YaHei';
                font-size: {size}px;
            }}
        """

        base_style_header = f"""
            QTextEdit {{
                border: none; 
                background-color: #0d1117; 
                font-size: {size}px;
                font-family: Consolas, 'Microsoft YaHei';
            }}
        """

        self.list_results.setStyleSheet(base_style_list)
        self.txt_detail.setStyleSheet(base_style_text)
        self.txt_header.setStyleSheet(base_style_header)

    def setup_shortcuts(self):
        self.shortcut_find = QShortcut(QKeySequence("Ctrl+F"), self)
        self.shortcut_find.activated.connect(self.focus_inner_search)

    def focus_inner_search(self):
        if self.isVisible() and hasattr(self, 'edit_inner_search'):
            self.edit_inner_search.setFocus()
            self.edit_inner_search.selectAll()

    def load_json(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择索引文件", "", "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            self._load_file(file_path)

    def refresh_current_file(self):
        if self.last_loaded_path and os.path.exists(self.last_loaded_path):
            self.edit_search.clear()
            self._load_file(self.last_loaded_path)
        else:
            QMessageBox.information(self, "提示", "尚未加载任何文件，或文件已不存在，无法刷新。")

    def _load_file(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                self.data = json.load(f)

            self.all_nodes = []
            root_nodes = self._smart_parse_structure(self.data)
            self._flatten_structure(root_nodes)

            if not self.all_nodes:
                self.txt_detail.setPlainText(
                    f"⚠️ 文件加载成功，但未解析到任何知识节点。\n"
                    f"文件: {os.path.basename(file_path)}\n"
                    f"请检查 JSON 是否包含 'structure' 或节点列表。"
                )
                self.list_results.clear()
                self.txt_header.clear()
                return

            self.last_loaded_path = file_path
            self.list_results.clear()
            for node in self.all_nodes:
                self._add_item_to_list(node)

            self.txt_detail.setPlainText(
                f"✅ 已成功加载索引文件\n"
                f"📄 文件: {os.path.basename(file_path)}\n"
                f"📊 共解析出 {len(self.all_nodes)} 个知识节点\n\n"
                f"请使用上方搜索框进行关键词召回，或点击左侧查看详情。"
            )
            self.txt_header.clear()
            self.edit_inner_search.clear()

        except Exception as e:
            import traceback
            error_msg = f"❌ 加载失败: {str(e)}\n\n{traceback.format_exc()}"
            self.txt_detail.setPlainText(error_msg)
            QMessageBox.critical(self, "错误", f"无法加载文件:\n{str(e)}")

    def _smart_parse_structure(self, data):
        if isinstance(data, list):
            return data
        elif isinstance(data, dict):
            if 'structure' in data and isinstance(data['structure'], list):
                return data['structure']
            if 'nodes' in data and isinstance(data['nodes'], list):
                return data['nodes']
            return [data]
        return []

    def _flatten_structure(self, nodes):
        if not nodes:
            return
        for item in nodes:
            if isinstance(item, dict):
                self.all_nodes.append(item)
                if 'nodes' in item and isinstance(item['nodes'], list):
                    self._flatten_structure(item['nodes'])

    def search_content(self):
        query = self.edit_search.text().strip().lower()
        self.list_results.clear()

        if not query:
            for node in self.all_nodes:
                self._add_item_to_list(node)
            self.txt_detail.setPlainText(f"显示全部 {len(self.all_nodes)} 个节点。")
            return

        results = 0
        for node in self.all_nodes:
            # 兼容两种格式的可搜索字段
            searchable = ' '.join([
                str(node.get('title', '')),
                str(node.get('metadata', {}).get('section_path', '')),
                str(node.get('text', '')),
                str(node.get('summary', '')),
                str(node.get('original_content', ''))
            ]).lower()

            if query in searchable:
                self._add_item_to_list(node)
                results += 1

        if results > 0:
            self.txt_detail.setPlainText(
                f"🔍 查询: \"{query}\"\n"
                f"✅ 找到 {results} 个匹配节点\n"
                f"请点击左侧列表查看详细内容。"
            )
        else:
            self.txt_detail.setPlainText(f"⚠️ 未找到包含 \"{query}\" 的内容。")

    def _add_item_to_list(self, node):
        # 兼容两种格式的标题提取
        title = node.get('title') or node.get('metadata', {}).get('section_path', '（无标题）')
        display = (title[:50] + '...') if len(title) > 50 else title
        item = QListWidgetItem(display)
        item.setToolTip(title)
        item.setData(Qt.UserRole, node)
        self.list_results.addItem(item)

    def display_node_detail(self, item):
        if item is None:
            return

        node = item.data(Qt.UserRole)
        if not node or not isinstance(node, dict):
            self.txt_header.clear()
            self.txt_detail.setPlainText("<i style='color:#8b949e;'>(无效节点数据)</i>")
            return

        # 兼容两种格式的标题
        title = node.get('title') or node.get('metadata', {}).get('section_path', '未命名章节')

        # 页码、node_id（RAG格式无页码）
        start = node.get('start_index', '-')
        end = node.get('end_index', '-')
        node_id = node.get('node_id', 'N/A')

        # 兼容两种格式的摘要和正文
        if 'original_content' in node:  # RAG格式
            summary = node.get('text', '')
            raw_text = node.get('original_content', '')
        else:  # 标准格式
            summary = node.get('summary', '')
            raw_text = node.get('text', '')

        header_html = f"""
        <h2 style='color: #58a6ff; margin: 0 0 10px 0;'>{html.escape(title)}</h2>
        <div style='background-color: #21262d; padding: 10px; border-radius: 6px; font-size: 0.9em;'>
            <span style='color: #8b949e; font-weight: bold;'>📄 物理页码:</span> 
            <span style='color: #c9d1d9;'>第 {start} - {end} 页</span>
            &nbsp;&nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;
            <span style='color: #8b949e; font-weight: bold;'>🆔 Node ID:</span> 
            <span style='color: #c9d1d9;'>{node_id}</span>
        </div>
        """
        if summary:
            header_html += f"""
            <div style='background-color: #1c2128; border-left: 4px solid #238636; padding: 10px; margin: 15px 0;'>
                <span style='color: #238636; font-weight: bold;'>💡 AI 摘要:</span><br>
                <span style='color: #c9d1d9;'>{html.escape(summary)}</span>
            </div>
            """
        self.txt_header.setHtml(header_html)

        if not raw_text:
            display_text = "<i style='color: #8b949e;'>(该节点无正文内容)</i>"
        else:
            display_text = html.escape(raw_text)

        self.txt_detail.setHtml(
            f"<div style='white-space: pre-wrap; font-family: Consolas, \"Microsoft YaHei\"; line-height: 1.7;'>{display_text}</div>"
        )

        QApplication.processEvents()
        self.highlight_text_in_detail()

    def highlight_text_in_detail(self):
        keyword = self.edit_inner_search.text().strip()
        if not keyword:
            return

        document = self.txt_detail.document()
        if document is None or document.isEmpty():
            return

        cursor = QTextCursor(document)
        cursor.select(QTextCursor.Document)
        clear_format = QTextCharFormat()
        clear_format.setBackground(Qt.transparent)
        clear_format.setForeground(QColor("#c9d1d9"))
        cursor.mergeCharFormat(clear_format)

        highlight_format = QTextCharFormat()
        highlight_format.setBackground(QColor("#d29922"))
        highlight_format.setForeground(QColor("black"))

        cursor = QTextCursor(document)
        cursor.setPosition(0)
        while True:
            cursor = document.find(keyword, cursor)
            if cursor.isNull():
                break
            cursor.mergeCharFormat(highlight_format)

    # ==================== 导出功能 ====================

    def export_all_nodes(self):
        if not self.all_nodes:
            QMessageBox.warning(self, "无数据", "当前未加载任何节点数据，无法导出。")
            return

        fmt = self.combo_export.currentText()
        ext_map = {
            "DOCX (Word)": ".docx",
            "TXT (纯文本)": ".txt",
            "CSV (表格)": ".csv",
            "XLSX (Excel)": ".xlsx"
        }
        default_ext = ext_map.get(fmt, ".txt")
        filter_map = {
            ".docx": "Word 文档 (*.docx)",
            ".txt": "文本文件 (*.txt)",
            ".csv": "CSV 文件 (*.csv)",
            ".xlsx": "Excel 文件 (*.xlsx)"
        }

        save_path, _ = QFileDialog.getSaveFileName(
            self, "导出知识节点", f"pageindex_export{default_ext}", filter_map[default_ext]
        )
        if not save_path:
            return

        try:
            if "DOCX" in fmt:
                self._export_docx(save_path)
            elif "TXT" in fmt:
                self._export_txt(save_path)
            elif "CSV" in fmt:
                self._export_csv(save_path)
            elif "XLSX" in fmt:
                self._export_xlsx(save_path)

            QMessageBox.information(self, "导出成功", f"已成功导出 {len(self.all_nodes)} 个节点至：\n{save_path}")

        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出时发生错误：\n{str(e)}")

    def _export_docx(self, path):
        if not HAS_DOCX:
            raise ImportError("未安装 python-docx，请运行: pip install python-docx")
        doc = Document()
        doc.add_heading("PageIndex 知识节点导出", 0)
        for node in self.all_nodes:
            title = node.get('title') or node.get('metadata', {}).get('section_path', '无标题')
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"页码: {node.get('start_index', '-')} - {node.get('end_index', '-')}")
            doc.add_paragraph(f"Node ID: {node.get('node_id', 'N/A')}")
            summary = node.get('summary') or node.get('text', '')
            if summary:
                p = doc.add_paragraph()
                p.add_run("AI 摘要: ").bold = True
                p.add_run(summary)
            raw_text = node.get('text', '') if 'original_content' not in node else node.get('original_content', '')
            doc.add_paragraph(raw_text or '(无正文)')
            doc.add_paragraph("-" * 40)
        doc.save(path)

    def _export_txt(self, path):
        with open(path, 'w', encoding='utf-8') as f:
            for i, node in enumerate(self.all_nodes, 1):
                title = node.get('title') or node.get('metadata', {}).get('section_path', '无标题')
                f.write(f"=== 节点 {i} ===\n")
                f.write(f"标题: {title}\n")
                f.write(f"页码: {node.get('start_index', '-')} - {node.get('end_index', '-')}\n")
                f.write(f"Node ID: {node.get('node_id', 'N/A')}\n")
                summary = node.get('summary') or node.get('text', '')
                if summary:
                    f.write(f"AI 摘要: {summary}\n")
                raw_text = node.get('text', '') if 'original_content' not in node else node.get('original_content', '')
                f.write(f"正文:\n{raw_text or '(无正文)'}\n")
                f.write("\n" + "-" * 60 + "\n\n")

    def _export_csv(self, path):
        if not HAS_PANDAS:
            raise ImportError("未安装 pandas，请运行: pip install pandas")
        data = []
        for node in self.all_nodes:
            title = node.get('title') or node.get('metadata', {}).get('section_path', '')
            summary = node.get('summary') or node.get('text', '')
            raw_text = node.get('text', '') if 'original_content' not in node else node.get('original_content', '')
            data.append({
                "Node ID": node.get('node_id', ''),
                "标题": title,
                "起始页": node.get('start_index', ''),
                "结束页": node.get('end_index', ''),
                "AI 摘要": summary,
                "正文内容": raw_text
            })
        pd.DataFrame(data).to_csv(path, index=False, encoding='utf-8-sig')

    def _export_xlsx(self, path):
        if not HAS_PANDAS:
            raise ImportError("未安装 pandas 和 openpyxl，请运行: pip install pandas openpyxl")
        data = []
        for node in self.all_nodes:
            title = node.get('title') or node.get('metadata', {}).get('section_path', '')
            summary = node.get('summary') or node.get('text', '')
            raw_text = node.get('text', '') if 'original_content' not in node else node.get('original_content', '')
            data.append({
                "Node ID": node.get('node_id', ''),
                "标题": title,
                "起始页": node.get('start_index', ''),
                "结束页": node.get('end_index', ''),
                "AI 摘要": summary,
                "正文内容": raw_text
            })
        pd.DataFrame(data).to_excel(path, index=False)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = PGIRecallWindow()
    window.show()
    sys.exit(app.exec_())